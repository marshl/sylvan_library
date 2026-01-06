"""
The module for the deck_export command
"""

import logging
import os
import re
from typing import List

import docx
from django.contrib.auth import get_user_model
from django.core.management.base import BaseCommand
from django.utils.text import slugify

from sylvan_library.cards.models.decks import Deck, DeckCard
from sylvan_library.cards.models.sets import Set
from website.templatetags.deck_templatetags import deck_card_group_count

logger = logging.getLogger("django")


class Command(BaseCommand):
    """
    The command for exporting user decks
    """

    help = "Exports user owned decks"

    def add_arguments(self, parser) -> None:
        # Positional arguments
        parser.add_argument(
            "username",
            nargs=1,
            type=str,
            help="The user whose decks should be exported",
        )

    def handle(self, *args, **options):
        output_directory = os.path.join("data_export", "output")
        username = options.get("username")[0]

        if not os.path.exists(output_directory):
            logger.error("Could not find directory %s", output_directory)
            return

        try:
            user = get_user_model().objects.get(username=username)
        except get_user_model().DoesNotExist:
            logger.error("User with name %s could not be found", username)
            return

        output_directory = os.path.join(output_directory, slugify(user.username))

        os.makedirs(output_directory, exist_ok=True)

        # self.export_decks(output_directory, user)
        self.export_pretty_decks(output_directory, user)

    def get_ordinal_suffix(self, n):
        """
        Convert an integer into its ordinal representation::

            make_ordinal(0)   => '0th'
            make_ordinal(3)   => '3rd'
            make_ordinal(122) => '122nd'
            make_ordinal(213) => '213th'
        """
        n = int(n)
        if 11 <= (n % 100) <= 13:
            suffix = "th"
        else:
            suffix = ["th", "st", "nd", "rd", "th"][min(n % 10, 4)]
        return suffix

    def export_pretty_decks(self, output_directory: str, user: get_user_model()):
        # deck = user.decks.order_by("-date_created").first()
        deck = Deck.objects.get(id=9624)

        template_document = docx.Document(
            os.path.join("data_export", "templates", "edh_deck_template.docx")
        )

        deck_release_date_str = deck.date_created.strftime("%d{} of %B, %Y").format(
            self.get_ordinal_suffix(deck.date_created.day)
        )

        latest_set = (
            Set.objects.filter(
                release_date__lte=deck.date_created,
                type__in=("masters", "expansion", "core"),
            )
            .order_by("-release_date")
            .first()
        )

        for paragraph in template_document.paragraphs:
            if "<DECK_NAME>" in paragraph.text:
                paragraph.text = paragraph.text.replace("<DECK_NAME>", deck.name)
            if "<DECK_SUBTITLE>" in paragraph.text and deck.subtitle:
                paragraph.text = paragraph.text.replace(
                    "<DECK_SUBTITLE>", deck.subtitle
                )
            if "<DECK_RELEASE_DATE>" in paragraph.text:
                paragraph.text = paragraph.text.replace(
                    "<DECK_RELEASE_DATE>", deck_release_date_str
                )
            if "<DECK_LATEST_SET>" in paragraph.text:
                paragraph.text = paragraph.text.replace(
                    "<DECK_LATEST_SET>", latest_set.name
                )

        card_groups = {group["code"]: group for group in deck.get_card_groups()}

        for table in template_document.tables:
            header = table.rows[0]
            is_basic_land = len(header.cells) == 2
            if is_basic_land:
                deck_cards = [
                    deck_card
                    for deck_card in card_groups["land"]["cards"]
                    if deck_card.card.faces.first()
                    .supertypes.filter(name="Basic")
                    .exists()
                ]
                group_name = "Basic Land"
            else:
                header_text = header.cells[0].text
                print(header_text)
                if header_text.endswith("General"):
                    card_group = card_groups["commander"]
                elif header_text.endswith("Nonbasic Land"):
                    card_group = card_groups["land"]
                else:
                    card_group = card_groups[header_text]

                deck_cards: List[DeckCard] = card_group["cards"]
                if header_text.endswith("Nonbasic Land"):
                    deck_cards = [
                        deck_card
                        for deck_card in deck_cards
                        if not deck_card.card.faces.first().supertypes.filter(
                            name="Basic"
                        )
                    ]
                    group_name = "Nonbasic Land"
                else:
                    group_name = card_group["name"]

            count = sum(deck_card.count for deck_card in deck_cards)
            header.cells[0].paragraphs[0].text = (
                f"{count}" if is_basic_land else f"{count} {group_name}"
            )

            if is_basic_land:
                header.cells[1].paragraphs[0].text = "Basic Land"

            if deck_cards:
                first_row = table.rows[1]
                if is_basic_land:
                    first_row.cells[0].paragraphs[0].text = str(deck_cards[0].count)
                first_row.cells[1 if is_basic_land else 0].paragraphs[
                    0
                ].text = deck_cards[0].card.name

                for deck_card in deck_cards[1:]:
                    row = table.add_row()
                    row.cells[0].paragraphs[0].style = (
                        first_row.cells[0].paragraphs[0].style
                    )
                    if is_basic_land:
                        row.cells[0].paragraphs[0].text = str(deck_card.count)
                        row.cells[1].paragraphs[0].style = (
                            first_row.cells[0].paragraphs[0].style
                        )
                        # row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    row.cells[1 if is_basic_land else 0].paragraphs[
                        0
                    ].text = deck_card.card.name

        deck_slug = slugify(f"{deck.date_created.isoformat()} {deck.name}")
        output_filename = os.path.join(output_directory, f"{deck_slug}.docx")
        logger.info("Saving output to %s", output_filename)
        template_document.save(output_filename)

    def export_decks(self, output_directory: str, user: get_user_model()):
        deck: Deck
        for deck in user.decks.filter(is_prototype=False):
            deck_slug = slugify(f"{deck.date_created.isoformat()} {deck.name}")
            filename = os.path.join(output_directory, f"{deck_slug}.txt")
            if os.path.exists(filename):
                continue
            print(deck.name, filename)

            deck = Deck.objects.prefetch_related("cards__card").get(id=deck.id)

            with open(filename, "w", encoding="utf-8") as file:
                file.write(f"// Name: {deck.name}\n")
                file.write(f"// Date: {deck.date_created.isoformat()}\n")
                if deck.description:
                    # Add // to newlines
                    commented_description = re.sub(
                        r"^(.+?)$", r"// \1", deck.description, flags=re.MULTILINE
                    )
                    # Remove only carriage returns
                    commented_description = re.sub(
                        r"\r", r"\n", commented_description, flags=re.MULTILINE
                    )
                    # Remove blank lines
                    commented_description = re.sub(
                        r"\n\s*\n", "\n", commented_description, flags=re.MULTILINE
                    )
                    file.write(f"// Description: \n{commented_description}\n")
                if deck.format:
                    pretty_format = next(
                        (f[1] for f in Deck.FORMAT_CHOICES if f[0] == deck.format),
                        deck.format,
                    )
                    file.write(f"// Format: {pretty_format}\n")

                file.write("\n")

                for card_group in deck.get_card_groups():
                    if not card_group["cards"]:
                        continue
                    file.write(
                        f"// {deck_card_group_count(card_group['cards'])} {card_group['name']}\n"
                    )
                    for card in card_group["cards"]:
                        file.write(f"{card.as_deck_text()}\n")

                sideboard = deck.get_cards("side")
                if sideboard:
                    file.write("\n")
                    file.write("// Sideboard:\n")
                    for card in sideboard:
                        file.write(f"{card.as_deck_text()}\n")
